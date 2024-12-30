
from PyPDF2 import PdfReader
from docx import Document

class Reader:
    def __init__(self, file):
        super().__init__()
        self.file = file

    def read_pdf(self):
        pdf_reader = PdfReader(self.file)
        text = ""
        for page in pdf_reader.pages:
            try:
                text += page.extract_text()
            except Exception as e:
                print('Error when reading pdf file:', e)
            
        return text

    def read_docx(self):
        #doc = Document(self.file)
        doc = Document()
        try:
            doc = Document(self.file)
            print("load doc success!")
        except Exception as e:
            print('Error when reading file:', e)
        text = ""
        for para in doc.paragraphs:
            text += para.text

        return text

    def read(self):
        if self.file.name.endswith(".pdf"):
            pdf_text = self.read_pdf()
            print("pdf text laf ........:" , pdf_text)
            return pdf_text
        
        elif self.file.name.endswith(".docx"):
            docx_text = self.read_docx()
            return docx_text
        
        return None